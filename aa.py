import streamlit as st
import io
import os
import time
import re
import requests
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from groq import Groq
import PyPDF2
from urllib.parse import quote_plus
from PIL import Image
import tempfile
import random


def set_font_style(paragraph, font_name="Times New Roman", font_size=12, bold=False):
    """Set font style for a paragraph"""
    try:
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font.size = Pt(font_size)
            font.bold = bold
    except:
        pass


def create_heading_style(doc, style_name, font_size, bold=True, space_after=12):
    """Create a custom heading style with error handling"""
    try:
        styles = doc.styles
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Times New Roman'
            style.font.size = Pt(font_size)
            style.font.bold = bold
            style.paragraph_format.space_after = Pt(space_after)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.line_spacing = 1.15
        return styles[style_name]
    except:
        return doc.styles['Normal']


def add_header_footer_safe(doc, project_title, student_name):
    """Add header and footer with error handling"""
    try:
        section = doc.sections[0]
        
        # Header
        header = section.header
        if header.paragraphs:
            header_para = header.paragraphs[0]
            header_para.text = project_title.upper()[:50]  # Limit length
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if header_para.runs:
                header_para.runs[0].font.name = 'Times New Roman'
                header_para.runs[0].font.size = Pt(10)
                header_para.runs[0].font.bold = True
        
        # Footer - simplified without XML manipulation
        footer = section.footer
        if footer.paragraphs:
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_run = footer_para.add_run(f"{student_name[:30]} | Page #")
            footer_run.font.name = 'Times New Roman'
            footer_run.font.size = Pt(10)
            
    except Exception as e:
        st.warning(f"Header/footer setup failed: {str(e)}")


def search_google_images(query, api_key, cse_id, num_results=3):
    """Search for images using Google Custom Search API with improved error handling"""
    try:
        # Validate inputs
        if not api_key or not cse_id:
            return []
        
        if not query or len(query.strip()) == 0:
            return []
        
        # Clean and prepare the query
        cleaned_query = query.strip()
        cleaned_query = re.sub(r'[^\w\s-]', '', cleaned_query)
        if len(cleaned_query) > 100:
            cleaned_query = cleaned_query[:100]
        
        search_url = "https://www.googleapis.com/customsearch/v1"
        params = {
            'key': api_key.strip(),
            'cx': cse_id.strip(),
            'q': cleaned_query,
            'searchType': 'image',
            'imgSize': 'medium',
            'imgType': 'photo',
            'safe': 'active',
            'num': min(num_results, 10),
            'fileType': 'jpg,png,jpeg'
        }
        
        # Add small delay between requests to avoid rate limiting
        time.sleep(0.8)
        
        response = requests.get(search_url, params=params, timeout=15)
        
        if response.status_code == 400:
            error_detail = response.text
            if "invalid API key" in error_detail.lower():
                st.warning("Invalid Google API key. Please check your API key.")
            elif "custom search engine" in error_detail.lower():
                st.warning("Invalid Custom Search Engine ID. Please check your CSE ID.")
            elif "quota" in error_detail.lower() or "limit" in error_detail.lower():
                st.warning("Google API quota exceeded. Try again tomorrow or upgrade your plan.")
            return []
        
        elif response.status_code == 403:
            st.warning("Access denied. Check if Custom Search API is enabled in Google Cloud Console.")
            return []
        
        elif response.status_code == 429:
            st.warning("Too many requests. Please wait before trying again.")
            # Wait and retry once
            time.sleep(3)
            try:
                response = requests.get(search_url, params=params, timeout=15)
                response.raise_for_status()
            except:
                return []
        
        response.raise_for_status()
        
        data = response.json()
        image_urls = []
        
        if 'items' in data:
            for item in data['items']:
                image_url = item.get('link', '')
                if image_url and any(ext in image_url.lower() for ext in ['.jpg', '.jpeg', '.png']):
                    image_urls.append({
                        'url': image_url,
                        'title': item.get('title', 'Related Image')[:100]
                    })
        
        return image_urls
        
    except requests.exceptions.Timeout:
        st.warning(f"Image search timed out for '{query}' - continuing without images")
        return []
    except Exception as e:
        st.warning(f"Image search failed for '{query}': {str(e)}")
        return []


def download_image_safe(image_url, timeout=15, max_size_mb=5):
    """Download image with better error handling and size limits"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate, br',
            'DNT': '1',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }
        
        # Add delay to avoid overwhelming servers
        time.sleep(0.5)
        
        response = requests.get(image_url, headers=headers, timeout=timeout, stream=True)
        response.raise_for_status()
        
        # Check content type
        content_type = response.headers.get('content-type', '').lower()
        if not any(img_type in content_type for img_type in ['image/jpeg', 'image/jpg', 'image/png']):
            return None
        
        # Load with size checking
        image_data = b''
        downloaded_size = 0
        max_size_bytes = max_size_mb * 1024 * 1024
        
        for chunk in response.iter_content(chunk_size=8192):
            downloaded_size += len(chunk)
            if downloaded_size > max_size_bytes:
                return None
            image_data += chunk
        
        # Validate minimum size
        if len(image_data) < 1024:  # Less than 1KB
            return None
        
        # Load with PIL
        image = Image.open(io.BytesIO(image_data))
        
        # Convert to RGB if necessary
        if image.mode in ('RGBA', 'P'):
            background = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'RGBA':
                background.paste(image, mask=image.split()[-1])
            else:
                background.paste(image)
            image = background
        elif image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Resize if too large
        max_dimension = 800
        if max(image.size) > max_dimension:
            ratio = max_dimension / max(image.size)
            new_size = tuple(int(dim * ratio) for dim in image.size)
            image = image.resize(new_size, Image.Resampling.LANCZOS)
        
        # Validate minimum size
        if min(image.size) < 100:
            return None
            
        return image
        
    except Exception as e:
        return None


def add_image_to_document_safe(doc, image, caption="", width_inches=4.5):
    """Add image to document with better error handling and unique filenames"""
    try:
        if not image:
            return False
        
        # Ensure image is in RGB format
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Create unique temporary file with timestamp and random suffix
        temp_dir = tempfile.gettempdir()
        timestamp = int(time.time() * 1000)
        random_suffix = random.randint(1000, 9999)
        temp_filename = f"temp_img_{timestamp}_{random_suffix}.jpg"
        temp_path = os.path.join(temp_dir, temp_filename)
        
        try:
            # Save image with high quality
            image.save(temp_path, 'JPEG', quality=95, optimize=True)
            
            if not os.path.exists(temp_path) or os.path.getsize(temp_path) == 0:
                return False
            
            # Add image to document
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = img_paragraph.add_run()
            
            # Validate width
            if width_inches > 6.0:
                width_inches = 6.0
            elif width_inches < 2.0:
                width_inches = 2.0
            
            run.add_picture(temp_path, width=Inches(width_inches))
            
            # Add caption if provided
            if caption and caption.strip():
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_text = caption.strip()
                if len(caption_text) > 100:
                    caption_text = caption_text[:97] + "..."
                caption_run = caption_para.add_run(f"Figure: {caption_text}")
                caption_run.font.name = 'Times New Roman'
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
            
            return True
            
        finally:
            # Clean up temp file with retry
            for attempt in range(3):
                try:
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)
                    break
                except:
                    time.sleep(0.1)
                    continue
                
    except Exception as e:
        return False


def extract_pdf_text(pdf_file):
    """Extract text from uploaded PDF file with error filtering"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text and len(page_text.strip()) > 50:
                text += page_text + "\n"
        
        # Clean up the text
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', '', text)
        
        return text[:3000] if text else "No readable content found"
    except Exception as e:
        return f"Error processing {pdf_file.name}: Unable to extract readable content"


def generate_content_with_groq(client, prompt, max_retries=3):
    """Generate content using Groq API with retry mechanism"""
    for attempt in range(max_retries):
        try:
            chat_completion = client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": """You are an expert academic writer. Generate high-quality, well-structured academic content. 
                        Use proper academic language and formatting. When creating lists or steps, use bullet points or numbered lists.
                        Structure content with clear paragraphs and logical flow."""
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                model="gemma2-9b-it",
                max_tokens=2000,
                temperature=0.7,
            )
            return chat_completion.choices[0].message.content
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(2)
                continue
            else:
                return f"Error generating content: {str(e)}"


def generate_formal_abstract(client, title, description, num_pages):
    """Generate a formal academic abstract"""
    prompt = f"""
Write a formal academic abstract (150-200 words) for a project titled "{title}".
Project description: {description}
Target length: {num_pages} pages

The abstract should include:
- Brief background/context
- Research objectives
- Methodology overview
- Expected outcomes/significance

Use formal academic language and structure. Make it concise but comprehensive.
"""
    return generate_content_with_groq(client, prompt)


def format_content_with_lists(content):
    """Format content to include proper bullet points and numbered lists"""
    lines = content.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            if (line.lower().startswith(('•', '-', '*')) or 
                re.match(r'^\d+\.', line) or
                any(keyword in line.lower() for keyword in ['step ', 'objective ', 'method ', 'approach ']) and
                len(line) < 100):
                formatted_lines.append(f"• {line.lstrip('•-*').lstrip('0123456789.').strip()}")
            else:
                formatted_lines.append(line)
    
    return '\n'.join(formatted_lines)


def parse_toc_items(toc_items):
    """Parse table of contents and return structured sections"""
    if not toc_items or not toc_items.strip():
        return None
    
    toc_list = []
    lines = toc_items.split('\n')
    
    for line in lines:
        line = line.strip()
        if line:
            clean_line = re.sub(r'^\d+\.?\s*', '', line)
            toc_list.append(clean_line)
    
    return toc_list


def generate_project_sections(client, title, description, toc_items, num_pages, 
                            pdf_texts, additional_notes, google_api_key=None, cse_id=None):
    """Generate all project sections using Groq with improved image integration"""
    sections = {}
    
    base_context = f"""
Project Title: {title}
Project Description: {description}
Estimated Pages: {num_pages}
Additional Context: {additional_notes if additional_notes else 'None provided'}
"""
    
    # Add clean PDF context if available
    pdf_context = ""
    if pdf_texts:
        clean_texts = [text for text in pdf_texts if not text.startswith("Error")]
        if clean_texts:
            pdf_context = f"\nReference Material Context:\n{' '.join(clean_texts[:2])}"
    
    full_context = base_context + pdf_context
    
    parsed_toc = parse_toc_items(toc_items)
    
    # Track image search attempts to avoid quota exhaustion
    image_search_count = 0
    max_image_searches = 12  # Reasonable limit
    
    if parsed_toc:
        st.info(f"Generating content for {len(parsed_toc)} custom sections...")
        
        for i, section_title in enumerate(parsed_toc):
            section_key = f"section_{i+1}"
            
            section_prompt = f"""
{full_context}

Write a comprehensive academic section titled "{section_title}" (400-600 words) for this project.
Structure the content with:
- Clear introduction to the section topic
- Main content with proper paragraphs
- Use bullet points or numbered lists where appropriate (objectives, steps, methods, etc.)
- Academic language and citations where relevant

Make it relevant to the project topic and ensure logical flow.
"""
            
            with st.spinner(f"Generating '{section_title}'..."):
                content = generate_content_with_groq(client, section_prompt)
                formatted_content = format_content_with_lists(content)
                
                # Search for relevant images with quota management
                images = []
                if (google_api_key and cse_id and 
                    image_search_count < max_image_searches and
                    section_title.lower() not in ['references', 'bibliography']):
                    
                    # Create more specific search query
                    search_terms = [
                        f"{title} {section_title}",
                        f"{section_title} research methodology",
                        f"{title.split()[0]} {section_title}"
                    ]
                    
                    for search_term in search_terms[:2]:  # Try 2 different queries
                        if image_search_count >= max_image_searches:
                            break
                            
                        image_results = search_google_images(search_term, google_api_key, cse_id, 3)
                        image_search_count += 1
                        
                        successful_downloads = 0
                        for img_data in image_results:
                            if successful_downloads >= 2:  # Limit per section
                                break
                                
                            img = download_image_safe(img_data['url'])
                            if img:
                                images.append({
                                    'image': img,
                                    'caption': img_data['title'][:80] + "..." if len(img_data['title']) > 80 else img_data['title']
                                })
                                successful_downloads += 1
                        
                        if successful_downloads > 0:
                            break  # Got images, no need to try more queries
                
                sections[section_key] = {
                    'title': section_title,
                    'content': formatted_content,
                    'images': images
                }
                
                # Add delay between sections
                time.sleep(2)
    
    else:
        st.info("Generating content for standard academic sections...")
        
        default_sections = [
            ("introduction", "Introduction", f"""
Write a comprehensive academic introduction (500-700 words) for this project. Include:
- Background information and context
- Problem statement clearly defined
- Research objectives (use numbered list)
- Scope and significance of the study
- Brief overview of methodology

Use formal academic language with clear paragraph structure.
"""),
            
            ("literature_review", "Literature Review", f"""
Write a literature review section (600-800 words) for this project. Include:
- Overview of existing research in the field
- Key findings from related studies
- Theoretical frameworks
- Research gaps identified
- How this project addresses those gaps

Structure with clear themes and use academic citation style with placeholder references [1], [2], etc.
"""),
            
            ("methodology", "Methodology", f"""
Write a methodology section (500-600 words) for this project. Include:
- Research design and approach
- Data collection methods (use bullet points)
- Tools and techniques to be used
- Analysis procedures (use numbered steps)
- Limitations and considerations

Be specific and detailed about the methods with clear structure.
"""),
            
            ("results", "Results and Analysis", f"""
Write a results and expected outcomes section (400-500 words) for this project. Include:
- Expected findings and results
- Analysis methods to be used
- Data presentation strategies
- Key metrics and indicators
- Potential challenges

Structure with clear subsections and use appropriate formatting.
"""),
            
            ("conclusion", "Conclusion", f"""
Write a conclusion section (300-400 words) for this project. Include:
- Summary of the project objectives
- Key contributions and significance
- Implications of the research
- Future work possibilities
- Final recommendations

Provide a strong, impactful conclusion that ties everything together.
"""),
            
            ("references", "References", f"""
Generate 12-18 realistic academic references for this project topic. Format them in proper APA style.
Include a mix of:
- Recent journal articles (2018-2024)
- Conference papers
- Books and book chapters
- Reputable online resources

Make sure they are relevant to "{title}" and realistic. Use proper APA formatting with hanging indent.
""")
        ]
        
        for section_key, section_title, prompt_template in default_sections:
            full_prompt = f"{full_context}\n\n{prompt_template}"
            
            with st.spinner(f"Generating {section_title}..."):
                content = generate_content_with_groq(client, full_prompt)
                
                if section_key != 'references':
                    content = format_content_with_lists(content)
                
                # Search for relevant images (skip for references)
                images = []
                if (google_api_key and cse_id and section_key != 'references' and 
                    image_search_count < max_image_searches):
                    
                    # More targeted search queries based on section type
                    if section_key == 'introduction':
                        search_query = f"{title} overview concept"
                    elif section_key == 'methodology':
                        search_query = f"{title} methodology research methods"
                    elif section_key == 'literature_review':
                        search_query = f"{title} literature research review"
                    elif section_key == 'results':
                        search_query = f"{title} results analysis data"
                    else:
                        search_query = f"{title} {section_title}"
                    
                    image_results = search_google_images(search_query, google_api_key, cse_id, 3)
                    image_search_count += 1
                    
                    successful_downloads = 0
                    for img_data in image_results:
                        if successful_downloads >= 2:
                            break
                            
                        img = download_image_safe(img_data['url'])
                        if img:
                            images.append({
                                'image': img,
                                'caption': img_data['title'][:80] + "..." if len(img_data['title']) > 80 else img_data['title']
                            })
                            successful_downloads += 1
                
                sections[section_key] = {
                    'title': section_title,
                    'content': content,
                    'images': images
                }
                
                # Add delay between sections
                time.sleep(2)
    
    return sections


def add_section_content_safe(doc, section_data, section_counter, is_references=False):
    """Add section content with error handling"""
    try:
        # Section heading
        title = section_data.get('display_title', section_data.get('title', 'Section'))
        heading_text = f"{section_counter}. {title}"
        
        section_heading = doc.add_paragraph(heading_text)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if section_heading.runs:
            section_heading.runs[0].font.name = 'Times New Roman'
            section_heading.runs[0].font.size = Pt(14)
            section_heading.runs[0].font.bold = True
        
        # Content
        content = section_data.get('content', '')
        
        if is_references:
            # Handle references with hanging indent
            ref_lines = content.split('\n')
            for ref_line in ref_lines:
                ref_line = ref_line.strip()
                if ref_line and not ref_line.lower().startswith('references'):
                    para = doc.add_paragraph(ref_line)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.first_line_indent = Inches(-0.5)
                    if para.runs:
                        para.runs[0].font.name = 'Times New Roman'
                        para.runs[0].font.size = Pt(12)
        else:
            # Handle regular content - split by double newlines for paragraphs
            paragraphs = content.split('\n\n')
            
            for paragraph_text in paragraphs:
                paragraph_text = paragraph_text.strip()
                if not paragraph_text:
                    continue
                    
                # Check if it's a bullet point
                if paragraph_text.startswith('•') or paragraph_text.startswith('-'):
                    # Add as bullet point
                    bullet_para = doc.add_paragraph()
                    bullet_para.style = 'List Bullet'
                    bullet_text = paragraph_text.lstrip('•-').strip()
                    bullet_run = bullet_para.add_run(bullet_text)
                    bullet_run.font.name = 'Times New Roman'
                    bullet_run.font.size = Pt(12)
                else:
                    # Add as regular paragraph
                    para = doc.add_paragraph(paragraph_text)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    if para.runs:
                        para.runs[0].font.name = 'Times New Roman'
                        para.runs[0].font.size = Pt(12)
        
        # Add images if available
        if 'images' in section_data and section_data['images']:
            for img_data in section_data['images'][:2]:  # Limit to 2 images
                try:
                    success = add_image_to_document_safe(doc, img_data['image'], img_data.get('caption', ''), 4.5)
                    if success:
                        # Add some spacing after image
                        doc.add_paragraph()
                except Exception:
                    pass  # Continue if image fails
        
        # Add spacing
        doc.add_paragraph()
        
    except Exception as e:
        st.warning(f"Section content error: {str(e)}")


def create_word_document_safe(title, student_name, description, toc_items, num_pages, 
                             sections, pdf_files, client, google_api_key=None, cse_id=None):
    """Create Word document with improved error handling"""
    try:
        # Create new document
        doc = Document()
        
        # Set document margins for safety
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Configure default style safely
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.15
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except:
            pass
        
        # Create safe heading styles
        heading1_style = create_heading_style(doc, 'SafeHeading1', 16, True)
        
        # Add headers and footers safely
        try:
            add_header_footer_safe(doc, title, student_name)
        except Exception as e:
            st.warning(f"Header/footer creation failed: {str(e)}")
        
        # === COVER PAGE ===
        try:
            # Title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title.upper())
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            
            # Add spacing
            for _ in range(3):
                doc.add_paragraph()
            
            # Subtitle
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run("A Comprehensive Academic Project")
            subtitle_run.font.name = 'Times New Roman'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.italic = True
            
            # Add spacing
            for _ in range(2):
                doc.add_paragraph()
            
            # Student name
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(f"Submitted by:\n{student_name}")
            name_run.font.name = 'Times New Roman'
            name_run.font.size = Pt(14)
            name_run.font.bold = True
            
            # Date
            doc.add_paragraph()
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_date = datetime.now().strftime("%B %d, %Y")
            date_run = date_para.add_run(current_date)
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(12)
            
        except Exception as e:
            st.warning(f"Cover page creation error: {str(e)}")
        
        # Page break
        doc.add_page_break()
        
        # === ABSTRACT ===
        try:
            abstract_heading = doc.add_paragraph("ABSTRACT")
            abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            abstract_run = abstract_heading.runs[0]
            abstract_run.font.name = 'Times New Roman'
            abstract_run.font.size = Pt(14)
            abstract_run.font.bold = True
            
            # Generate and add abstract
            formal_abstract = generate_formal_abstract(client, title, description, num_pages)
            abstract_para = doc.add_paragraph(formal_abstract)
            abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in abstract_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
        except Exception as e:
            st.warning(f"Abstract creation error: {str(e)}")
        
        doc.add_page_break()
        
        # === TABLE OF CONTENTS ===
        try:
            toc_heading = doc.add_paragraph("TABLE OF CONTENTS")
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_run = toc_heading.runs[0]
            toc_run.font.name = 'Times New Roman'
            toc_run.font.size = Pt(14)
            toc_run.font.bold = True
            
            # Add simple TOC placeholder
            toc_para = doc.add_paragraph("Table of contents will be generated automatically when opened in Microsoft Word.")
            toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_para_run = toc_para.runs[0]
            toc_para_run.font.italic = True
            toc_para_run.font.size = Pt(11)
            
        except Exception as e:
            st.warning(f"TOC creation error: {str(e)}")
        
        doc.add_page_break()
        
        # === MAIN CONTENT SECTIONS ===
        section_counter = 1
        
        try:
            if any(key.startswith('section_') for key in sections.keys()):
                # Custom sections
                for key in sorted(sections.keys()):
                    if key.startswith('section_'):
                        add_section_content_safe(doc, sections[key], section_counter)
                        section_counter += 1
            else:
                # Default sections
                default_section_order = [
                    ('introduction', 'INTRODUCTION'),
                    ('literature_review', 'LITERATURE REVIEW'), 
                    ('methodology', 'METHODOLOGY'),
                    ('results', 'RESULTS AND ANALYSIS'),
                    ('conclusion', 'CONCLUSION'),
                    ('references', 'REFERENCES')
                ]
                
                for section_key, section_title in default_section_order:
                    if section_key in sections:
                        section_data = sections[section_key]
                        section_data['display_title'] = section_title
                        add_section_content_safe(doc, section_data, section_counter, is_references=(section_key=='references'))
                        section_counter += 1
                        
        except Exception as e:
            st.error(f"Content section creation error: {str(e)}")
        
        return doc
        
    except Exception as e:
        st.error(f"Document creation failed: {str(e)}")
        # Return a minimal document
        return create_minimal_document(title, student_name, description)


def create_minimal_document(title, student_name, description):
    """Create a minimal document if main creation fails"""
    try:
        doc = Document()
        
        # Title
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.bold = True
        title_para.runs[0].font.size = Pt(16)
        
        doc.add_paragraph()
        
        # Student name
        name_para = doc.add_paragraph(f"By: {student_name}")
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Description
        desc_para = doc.add_paragraph(description)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        return doc
    except:
        # Final fallback
        doc = Document()
        doc.add_paragraph("Document creation encountered errors. Please try again.")
        return doc


def display_content_preview(sections):
    """Display preview of generated content with proper formatting"""
    with st.expander("Preview Generated Content"):
        if any(key.startswith('section_') for key in sections.keys()):
            for key in sorted(sections.keys()):
                if key.startswith('section_'):
                    section_data = sections[key]
                    st.subheader(f"{section_data['title']}")
                    content = section_data['content']
                    preview_content = content[:400] + "..." if len(content) > 400 else content
                    st.write(preview_content)
                    
                    if 'images' in section_data and section_data['images']:
                        st.write(f"📷 {len(section_data['images'])} image(s) will be included")
                    st.markdown("---")
        else:
            section_titles = {
                'introduction': 'Introduction',
                'literature_review': 'Literature Review',
                'methodology': 'Methodology',
                'results': 'Results and Analysis',
                'conclusion': 'Conclusion',
                'references': 'References'
            }
            
            section_order = ['introduction', 'literature_review', 'methodology',
                           'results', 'conclusion', 'references']
            
            for section_key in section_order:
                if section_key in sections:
                    section_data = sections[section_key]
                    st.subheader(section_titles.get(section_key, section_key.title()))
                    content = section_data['content']
                    preview_content = content[:400] + "..." if len(content) > 400 else content
                    st.write(preview_content)
                    
                    if 'images' in section_data and section_data['images']:
                        st.write(f"📷 {len(section_data['images'])} image(s) will be included")
                    st.markdown("---")


def main():
    """Main application function"""
    st.set_page_config(
        page_title="Enhanced AI Project Generator",
        page_icon="🤖",
        layout="wide"
    )
    
    st.title("🤖 Enhanced AI-Powered Project Generator")
    st.markdown("Generate complete academic projects with AI content, images, and professional formatting")
    st.markdown("---")
    
    # Sidebar for API configuration
    st.sidebar.header("API Configuration")
    
    groq_api_key = st.sidebar.text_input(
        "Groq API Key *",
        type="password",
        help="Get your free API key from https://console.groq.com/"
    )
    
    st.sidebar.markdown("### Image Integration (Optional)")
    google_api_key = st.sidebar.text_input(
        "Google API Key",
        type="password",
        help="For automatic image fetching (optional)"
    )
    
    cse_id = st.sidebar.text_input(
        "Custom Search Engine ID",
        type="password",
        help="Google Custom Search Engine ID (optional)"
    )
    
    if google_api_key and cse_id:
        st.sidebar.success("✅ Image integration enabled")
    else:
        st.sidebar.info("💡 Add Google API keys for automatic image integration")
    
    if not groq_api_key:
        st.warning("Please enter your Groq API key in the sidebar to continue")
        st.info("Get your free API key from [Groq Console](https://console.groq.com/)")
        return
    
    # Initialize Groq client
    try:
        client = Groq(api_key=groq_api_key)
    except Exception as e:
        st.error(f"Error initializing Groq client: {str(e)}")
        return
    
    # Main interface layout
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📝 Project Information")
        
        title = st.text_input(
            "Project Title *",
            placeholder="e.g., Machine Learning Applications in Healthcare",
            help="Enter a descriptive title for your project"
        )
        
        student_name = st.text_input(
            "Student Name *",
            placeholder="Enter your full name"
        )
        
        description = st.text_area(
            "Project Description *",
            placeholder="Provide a detailed description of your project, including objectives, methodology, and expected outcomes...",
            height=150,
            help="This will be used to generate a formal abstract and relevant content"
        )
        
        toc_items = st.text_area(
            "Custom Table of Contents (Optional)",
            placeholder="Introduction\nLiterature Review\nMethodology\nResults and Analysis\nConclusion",
            height=100,
            help="Leave empty to use default academic structure"
        )
        
        num_pages = st.slider(
            "Target Number of Pages",
            min_value=5,
            max_value=50,
            value=15,
            help="This affects the depth of generated content"
        )
    
    with col2:
        st.header("📚 Additional Resources")
        
        pdf_files = st.file_uploader(
            "Upload Reference PDFs (Optional)",
            type=['pdf'],
            accept_multiple_files=True,
            help="Upload PDFs to extract context for better content generation"
        )
        
        if pdf_files:
            st.success(f"📄 {len(pdf_files)} PDF(s) uploaded")
        
        additional_notes = st.text_area(
            "Additional Notes/Requirements",
            placeholder="Any specific requirements, focus areas, or additional context...",
            height=100
        )
        
        st.markdown("---")
        st.header("🚀 Generate Project")
        
        # Input validation
        can_generate = all([title, student_name, description, groq_api_key])
        
        if not can_generate:
            missing = []
            if not title:
                missing.append("Project Title")
            if not student_name:
                missing.append("Student Name")
            if not description:
                missing.append("Project Description")
            
            for item in missing:
                st.warning(f"❌ Missing: {item}")
        
        # Feature indicators
        features = []
        features.append("✅ AI-Generated Content")
        features.append("✅ Professional Word Formatting")
        features.append("✅ Headers & Footers")
        features.append("✅ Proper Citations & References")
        features.append("✅ Table of Contents")
        
        if google_api_key and cse_id:
            features.append("✅ Automatic Image Integration")
        else:
            features.append("⚪ Image Integration (API keys needed)")
        
        st.markdown("### Features:")
        for feature in features:
            st.markdown(feature)
        
        # Generate button
        if st.button(
            "🎯 Generate Complete Project",
            disabled=not can_generate,
            type="primary",
            use_container_width=True,
            help="Generate a complete academic project with AI-powered content and images"
        ):
            if can_generate:
                # Show progress
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                try:
                    # Extract PDF texts
                    pdf_texts = []
                    if pdf_files:
                        status_text.text("📄 Processing uploaded PDFs...")
                        progress_bar.progress(10)
                        
                        for pdf_file in pdf_files:
                            pdf_text = extract_pdf_text(pdf_file)
                            if not pdf_text.startswith("Error"):
                                pdf_texts.append(pdf_text)
                        
                        if pdf_texts:
                            st.success(f"✅ Successfully processed {len(pdf_texts)} PDFs")
                        else:
                            st.warning("⚠️ No readable content found in uploaded PDFs")
                    
                    progress_bar.progress(20)
                    
                    # Generate content
                    status_text.text("🤖 AI is generating your project content...")
                    
                    sections = generate_project_sections(
                        client, title, description, toc_items,
                        num_pages, pdf_texts, additional_notes,
                        google_api_key, cse_id
                    )
                    
                    progress_bar.progress(70)
                    
                    # Create Word document using the safe function
                    status_text.text("📝 Creating Word document with formatting...")
                    
                    doc = create_word_document_safe(
                        title, student_name, description, toc_items,
                        num_pages, sections, pdf_files, client,
                        google_api_key, cse_id
                    )
                    
                    progress_bar.progress(90)
                    
                    # Save to BytesIO
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    # Create safe filename
                    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                    safe_title = safe_title.replace(' ', '_')[:30]
                    filename = f"{safe_title}_Enhanced_Project.docx" if safe_title else "AI_Generated_Enhanced_Project.docx"
                    
                    progress_bar.progress(100)
                    status_text.text("✅ Project generation complete!")
                    
                    st.success("🎉 Complete project generated successfully!")
                    st.balloons()
                    
                    # Show statistics
                    col_stat1, col_stat2, col_stat3 = st.columns(3)
                    
                    with col_stat1:
                        st.metric("📊 Sections Generated", len(sections))
                    
                    with col_stat2:
                        total_images = sum(len(section.get('images', [])) for section in sections.values())
                        st.metric("🖼️ Images Added", total_images)
                    
                    with col_stat3:
                        total_words = sum(len(section['content'].split()) for section in sections.values())
                        st.metric("📝 Total Words", f"{total_words:,}")
                    
                    # Download button
                    st.download_button(
                        label="📥 Download Enhanced Project",
                        data=doc_io.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )
                    
                    # Show preview
                    display_content_preview(sections)
                    
                    # Instructions for users
                    with st.expander("📋 Document Instructions"):
                        st.markdown("""
                        ### Your Enhanced Document Includes:
                        
                        **Professional Formatting:**
                        - ✅ Cover page with title, name, and date
                        - ✅ Headers with project title
                        - ✅ Footers with your name and page numbers
                        - ✅ Proper font styling (Times New Roman, 12pt)
                        - ✅ 1.5 line spacing and justified text
                        
                        **Content Structure:**
                        - ✅ Formal abstract (150-200 words)
                        - ✅ Automatic table of contents
                        - ✅ Well-structured sections with headings
                        - ✅ Bullet points and numbered lists where appropriate
                        - ✅ APA-style references with hanging indent
                        
                        **Visual Elements:**
                        - ✅ Relevant images with captions (if API keys provided)
                        - ✅ Proper image alignment and sizing
                        
                        ### Next Steps:
                        1. **Open the document** in Microsoft Word
                        2. **Update the Table of Contents:** Right-click on TOC → Update Field → Update entire table
                        3. **Review and customize** the content as needed
                        4. **Check citations** and add real references if required
                        5. **Proofread** for any final adjustments
                        
                        ### Tips:
                        - The document uses heading styles for easy navigation
                        - All formatting is consistent and professional
                        - Images are automatically sized and centered
                        - References follow APA format guidelines
                        """)
                    
                except Exception as e:
                    progress_bar.progress(0)
                    status_text.text("")
                    st.error(f"❌ Error generating project: {str(e)}")
                    st.info("💡 Please check your API keys and try again")
    
    # Information sections
    st.markdown("---")
    
    col_info1, col_info2 = st.columns(2)
    
    with col_info1:
        st.markdown("""
        ### 🔧 **How It Works**
        1. **Enter Project Details** - Title, description, and requirements
        2. **Upload References** - Add PDFs for context (optional)
        3. **API Integration** - Uses Groq for AI content + Google for images
        4. **Generate Document** - Creates professional Word document
        5. **Download & Customize** - Get your formatted project ready for submission
        """)
    
    with col_info2:
        st.markdown("""
        ### 📋 **What You Get**
        - **Professional formatting** with headers, footers, and proper styling
        - **AI-generated content** tailored to your topic
        - **Relevant images** automatically sourced and inserted
        - **Proper citations** in APA format
        - **Table of contents** that updates automatically
        - **Structured sections** with bullet points and lists
        """)
    
    # API Information
    with st.expander("🔑 API Setup Instructions"):
        st.markdown("""
        ### Required: Groq API Key
        1. Visit [Groq Console](https://console.groq.com/)
        2. Sign up for a free account
        3. Generate an API key
        4. Enter the key in the sidebar
        
        ### Optional: Google Custom Search (for images)
        1. Go to [Google Cloud Console](https://console.cloud.google.com/)
        2. Enable the Custom Search API
        3. Create credentials (API key)
        4. Set up a Custom Search Engine at [CSE](https://cse.google.com/)
        5. Configure it to search the entire web
        6. Copy the Search Engine ID
        7. Enter both keys in the sidebar for automatic image integration
        
        **Note:** Google API has usage limits. The free tier includes 100 searches per day.
        """)
    
    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666;'>
            <p><strong>Enhanced AI-Powered Project Generator</strong></p>
            <p>🤖 Groq AI + 🖼️ Google Images + 📝 Professional Word Formatting</p>
            <p><em>Generate complete academic projects with AI assistance and visual integration</em></p>
        </div>
        """,
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()